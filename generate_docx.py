import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def find_comments_in_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    results = []

    # Find multiline comments /* ... */
    multiline_matches = re.finditer(r'/\*(.*?)\*/', content, flags=re.DOTALL)
    for m in multiline_matches:
        text = m.group(1)
        if "MAPA DO TESOURO" in text:
            clean_text = "\n".join(line.strip(' *') for line in text.strip().split('\n'))
            results.append(clean_text)

    # Find single line comments // ...
    # We will split by lines and look for contiguous // blocks
    lines = content.split('\n')
    in_block = False
    current_block = []
    has_treasure_map = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('//'):
            current_block.append(stripped[2:].strip())
            if "MAPA DO TESOURO" in stripped:
                has_treasure_map = True
            in_block = True
        else:
            if in_block:
                if has_treasure_map:
                    results.append("\n".join(current_block))
                current_block = []
                has_treasure_map = False
                in_block = False
                
    if in_block and has_treasure_map:
        results.append("\n".join(current_block))

    return results

def main():
    doc = Document()
    doc.add_heading('Especificações do Sistema (Mapa do Tesouro)', 0)

    p = doc.add_paragraph('Este documento contém todas as especificações e documentações extraídas dos comentários "Mapa do Tesouro" no código fonte.')
    
    src_dir = 'src'
    
    for root, dirs, files in os.walk(src_dir):
        for file in files:
            if file.endswith(('.ts', '.tsx')):
                filepath = os.path.join(root, file)
                comments = find_comments_in_file(filepath)
                if comments:
                    doc.add_heading(f'Arquivo: {filepath}', level=1)
                    for comment in comments:
                        p = doc.add_paragraph(comment)
                        p.style.font.name = 'Courier New'
                        p.style.font.size = Pt(10)

    output_path = 'Especificacoes_Mapa_do_Tesouro.docx'
    doc.save(output_path)
    print(f"Docx saved to {output_path}")

if __name__ == '__main__':
    main()
